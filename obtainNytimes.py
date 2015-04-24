import urllib.request, urllib.error, sys, re
from http.cookiejar import CookieJar
from docx import Document


def writeDoc(title,content):
	docName = title+'.docx'
	document = Document()

	document.add_heading(title, 0)

	document.add_paragraph(content)
	
	document.save(docName)

def ObtainContent(pageContent):
	#obtain title
	title = ''
	for ln in pageContent:
		#print(ln)
		mat = re.search(b'<h1 itemprop="headline" id="story-heading" class="story-heading">.*</h1>', ln)
		if mat:
			headline = mat.group(0).decode('utf-8')
			length = len(headline)
			i = 0
			while i < length:
				if headline[i] == '<':
				#find >
					z = i + 1
					while z < length and headline[z] != '>':
						z = z + 1
					i = z + 1
					while i < length and headline[i] != '<':
						title = title + headline[i]	
						i = i + 1
				else:
					i = i + 1
			break

	#obtain content	
	#step 1: get all content with label p	
	paraList = []
	for ln in pageContent:
		mat = re.findall(b'<p class="story-body-text story-content".*?</p>', ln)
		for m in mat:
			paraList.append(m)

	#step 2: fetch content between <p> </p>
	para = ''
	for e in paraList:
		extract = e.decode('utf-8')
		length = len(extract)
		i = 0
		while i < length:
			if extract[i] == '<':
			#find >
				z = i + 1
				while z < length and extract[z] != '>':
					z = z + 1
				i = z + 1
				while i < length and extract[i] != '<':
					para = para + extract[i]	
					i = i + 1
			else:
				i = i + 1
		para = para + '\n'

	return (title,para)
	
def fetchWebPages(website):
	cj = CookieJar()
	
	opener = urllib.request.build_opener(urllib.request.HTTPCookieProcessor(cj)) 
	
	page = None
	try:
		page = opener.open(website)
	except urllib.error.HTTPError as e:
		print("HTTP Error: " + str(e.code))
	except urllib.error.URLError as e:
		print("URL Error: " + str(e.reason))
	except Exception as e:
		print(e)
	return page

def ObtainNyTimes():

	page = fetchWebPages(sys.argv[1])
	if page:
		(title, paras) = ObtainContent(page)

		if title == '' or paras == '':

			print("Your Newspaper doesn't have the key elements, Please make sure that your url is Nytimes....")

		else:

			writeDoc(title,paras)

			print("Fetch Your Newspaper Successfully..........")
	else:

		print("Failed to Fetch your Newspaper, Please Check your internet connection.....")
	

if __name__ == "__main__":
 	ObtainNyTimes()
